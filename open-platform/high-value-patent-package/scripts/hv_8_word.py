# -*- coding: utf-8 -*-
"""Generate Word report from final_records.json. Matches HTML substance."""
import io, datetime, requests
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from hv_common import jload, load_query

BLUE = RGBColor(0x2f, 0x6f, 0xed)
DARK = RGBColor(0x33, 0x33, 0x33)

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true')
    trPr.append(th)

def shade_runs(cell, size=8, bold=False, color=None, white=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(size)
            r.font.bold = bold
            if white:
                r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            elif color:
                r.font.color.rgb = color

def fetch_img(url, timeout=30):
    try:
        r = requests.get(url, timeout=timeout)
        if r.status_code == 200 and r.content[:4] in (b'\x89PNG', b'\xff\xd8\xff\xe0', b'\xff\xd8\xff\xe1', b'GIF8'):
            return io.BytesIO(r.content)
    except Exception:
        return None
    return None

def add_hyperlink(cell, url, text, size=7):
    """Insert a clickable external hyperlink into a table cell."""
    para = cell.paragraphs[0]
    part = para.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '2F6FED'); rPr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    para._p.append(hyperlink)

def add_heading(doc, text, size, color=DARK, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = True; r.font.color.rgb = color
    return p

def add_kv_table(doc, rows, widths):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            shade_runs(cells[ci], size=9, bold=(ri == 0), white=(ri == 0))
            if ri == 0:
                set_cell_bg(cells[ci], '42506B')
    return t

def build_front(doc, m, sel):
    now = datetime.datetime.now().strftime('%Y-%m-%d %H:%M')
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('高价值专利包筛选报告'); r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = BLUE
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(f'生成时间：{now}　|　数据来源：PatSnap / 智慧芽 Connect API（实时检索，未人工编造）')
    rs.font.size = Pt(9); rs.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    add_heading(doc, '一、检索与筛选概览', 14, BLUE)
    p = doc.add_paragraph()
    p.add_run(f"本次检索返回 {m['candidate_count']} 件候选专利，筛选 {m['selected_count']} 件高价值专利，"
              f"占 {m['ratio']}%（10% 下限 = {m['candidate_count']*10//100} 件，15% 上限 = "
              f"{m['sel_max_15pct']} 件；按默认 10% 向上取整选取）。入选评分区间 "
              f"{sel[-1]['score']} ~ {sel[0]['score']}。").font.size = Pt(10.5)
    add_kv_table(doc, [
        ['指标', '数值'],
        ['检索返回候选专利', m['candidate_count']],
        ['筛选高价值专利', m['selected_count']],
        ['高价值占比', f"{m['ratio']}%"],
        ['最高评分 / 入选最低评分', f"{sel[0]['score']} / {sel[-1]['score']}"],
    ], None)
    add_heading(doc, '检索式', 11)
    q = doc.add_paragraph(); qr = q.add_run(load_query())
    qr.font.size = Pt(8); qr.font.name = 'Consolas'

def build_method(doc, m):
    add_heading(doc, '二、评分标准与方法', 14, BLUE)
    add_kv_table(doc, [
        ['指标', '权重', '数据来源 / 口径'],
        ['简单同族被引专利数量多', '30%', 'P015 cited_by_simple_family，候选集百分位×30'],
        ['简单同族专利数量多', '30%', 'P014 len(simple_family)，候选集百分位×30'],
        ['属于核心发明人专利', '20%', '命中候选集发明人专利数 Top5 得 20，否则 0'],
        ['出现过专利法律事件', '20%', 'P034诉讼/P027无效复审/P028许可/P029权利转移，任一非空得20'],
    ], None)
    add_heading(doc, '核心发明人（候选集专利数 Top5）', 11)
    rows = [['排名', '发明人', '候选集专利数']]
    for i, t in enumerate(m['top5_inventors']):
        rows.append([i + 1, t['name'], t['count']])
    add_kv_table(doc, rows, None)
    note = doc.add_paragraph()
    nr = note.add_run('口径说明：智慧芽发明人字段为「姓, 名|姓, 名」格式，逗号是单个发明人姓名内的姓/名分隔符，'
        '竖线「|」才是发明人之间的分隔符。为避免把「TANNER, CHRISTOPHER RICHARD」拆成两个人，本次仅按'
        '「|、；、换行」切分发明人（与通用规则中“同时按逗号切分”有意偏离）。中英文同族对同一发明人会以不同'
        '语言分别出现，跨语言不做归并，核心发明人基于检索返回候选集原样统计。')
    nr.font.size = Pt(8.5); nr.font.color.rgb = RGBColor(0x80, 0x60, 0x00)

COLS = ['排名', '评分', '被选定为高价值的原因', '公开公告号', '标题', '摘要附图',
        '[标]当前申请(专利权)人', '简单法律状态', 'Patsnap专利标题', 'AI技术问题',
        'AI技术手段', 'AI技术功效', '同族被引', '同族数', '核心发明人', '法律事件', '数据缺口']
WIDTHS = [0.4, 0.5, 2.2, 1.1, 1.6, 1.1, 1.3, 0.9, 1.8, 2.0, 2.0, 2.0, 0.6, 0.6, 1.1, 1.0, 1.2]

def ev_text(r):
    ev = r['legal_events']
    if not ev:
        return '无'
    return '、'.join(f"{k}{('×'+str(v)) if v>1 else ''}" for k, v in ev.items())

def core_text(r):
    return ('是：' + '、'.join(r['matched_inventors'])) if r['core_inventor'] else '否'

def build_table(doc, sel, embed_img=True):
    add_heading(doc, f'三、高价值专利清单（{len(sel)} 件）', 14, BLUE)
    hint = doc.add_paragraph()
    hr = hint.add_run('提示：点击「公开公告号」可在智慧芽 PatSnap 数据库中打开该专利详情页（需登录智慧芽账号）。')
    hr.font.size = Pt(8.5); hr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    t = doc.add_table(rows=1, cols=len(COLS)); t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, c in enumerate(COLS):
        hdr[i].text = c; shade_runs(hdr[i], size=7.5, bold=True, white=True)
        set_cell_bg(hdr[i], '2F6FED')
    set_repeat_header(t.rows[0])
    for i, w in enumerate(WIDTHS):
        for cell in t.columns[i].cells:
            cell.width = Inches(w)
    for r in sel:
        cells = t.add_row().cells
        vals = [r['rank'], r['score'], r['rationale'], r['pn'], r['title'], '',
                r['current_assignee'], r['legal_status'], r['patsnap_title'],
                r['tech_problem'], r['tech_approach'], r['benefit'],
                (r['cited_by_simple_family'] if not r['cited_missing'] else '未获取'),
                r['simple_family_count'], core_text(r), ev_text(r),
                '、'.join(r['gaps']) if r['gaps'] else '无']
        for ci, v in enumerate(vals):
            if ci in (3, 5):
                continue
            cells[ci].text = str(v); shade_runs(cells[ci], size=7)
        # publication number cell -> hyperlink to Zhihuiya patent view
        url = r.get('view_url')
        if url:
            add_hyperlink(cells[3], url, r['pn'], size=7)
        else:
            cells[3].text = str(r['pn']); shade_runs(cells[3], size=7)
        # drawing cell
        d = r['drawing']
        if embed_img and isinstance(d, str) and d.startswith('http'):
            bio = fetch_img(d)
            if bio:
                try:
                    cells[5].paragraphs[0].add_run().add_picture(bio, width=Inches(0.95))
                except Exception:
                    cells[5].text = '附图嵌入失败'; shade_runs(cells[5], size=7)
            else:
                cells[5].text = '附图链接已失效'; shade_runs(cells[5], size=7)
        else:
            cells[5].text = str(d); shade_runs(cells[5], size=7)
    for i, w in enumerate(WIDTHS):
        for cell in t.columns[i].cells:
            cell.width = Inches(w)

def build_gaps(doc, sel):
    from collections import Counter
    gc = Counter()
    for r in sel:
        for g in r['gaps']:
            gc[g] += 1
    add_heading(doc, '四、数据缺口与接口情况', 14, BLUE)
    if gc:
        for k, v in gc.most_common():
            p = doc.add_paragraph(style='List Bullet'); p.add_run(f'{k}：{v} 件').font.size = Pt(9)
    else:
        doc.add_paragraph('无')
    p = doc.add_paragraph()
    p.add_run('说明：缺失字段标记为「未获取」并保留在追溯数据 高价值专利包筛选数据.json 中；法律事件仅作为价值信号，'
        '不构成专利稳定性或可执行性的法律结论。本次候选集法律事件以「权利转移」为主（304 件），许可 1 件，'
        '未检索到诉讼与无效/复审记录。').font.size = Pt(9)

def set_landscape(doc):
    sec = doc.sections[0]
    from docx.enum.section import WD_ORIENT
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Cm(42), Cm(29.7)
    for s in (sec,):
        s.left_margin = Cm(1.2); s.right_margin = Cm(1.2)
        s.top_margin = Cm(1.2); s.bottom_margin = Cm(1.2)

if __name__ == '__main__':
    import sys
    embed = '--noimg' not in sys.argv
    f = jload('final_records.json')
    m, sel = f['meta'], f['selected']
    doc = Document()
    set_landscape(doc)
    style = doc.styles['Normal']; style.font.name = '微软雅黑'; style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:eastAsia'), '微软雅黑')
    build_front(doc, m, sel)
    build_method(doc, m)
    build_table(doc, sel, embed_img=embed)
    build_gaps(doc, sel)
    doc.save('高价值专利包筛选报告.docx')
    print('Word saved; embed_img=', embed, 'rows=', len(sel))
