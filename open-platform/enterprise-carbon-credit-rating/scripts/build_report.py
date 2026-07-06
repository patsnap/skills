#!/usr/bin/env python3
"""Build a polished enterprise carbon credit rating DOCX from rating_result.json."""

from __future__ import annotations

import argparse
import json
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


FONT_EN = "Arial"
FONT_CN = "微软雅黑"

NAVY = "0B2545"
TEAL = "1F6F6D"
GREEN = "2E7D32"
GOLD = "B68B2E"
RED = "9B1C1C"
SLATE = "455A64"
INK = "1F2933"
MUTED = "6B7780"
BORDER = "D9E1E5"
MIST = "F4F7F8"
PALE_GREEN = "E9F5EF"
PALE_GOLD = "FFF6E3"
PALE_RED = "FDECEC"
PALE_BLUE = "EAF2F6"
WHITE = "FFFFFF"


def rgb(hex_value: str) -> RGBColor:
    value = hex_value.strip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def fmt(value: Any, suffix: str = "") -> str:
    if value in (None, "", []):
        return "未采集数据"
    if isinstance(value, float):
        return f"{value:.2f}{suffix}"
    return f"{value}{suffix}"


def pct(value: Any) -> str:
    if value in (None, ""):
        return "未采集数据"
    return f"{float(value) * 100:.0f}%"


def set_run_font(
    run,
    *,
    size: float = 10.5,
    color: str = INK,
    bold: bool = False,
    italic: bool = False,
    name: str = FONT_EN,
):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)


def para_text(
    parent,
    text: str = "",
    *,
    size: float = 10.5,
    color: str = INK,
    bold: bool = False,
    italic: bool = False,
    align: int | None = None,
    before: float = 0,
    after: float = 6,
    line_spacing: float = 1.12,
):
    para = parent.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = line_spacing
    if align is not None:
        para.alignment = align
    if text:
        run = para.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return para


def paragraph_border_bottom(paragraph, color: str = BORDER, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, bottom: int = 100, start: int = 140, end: int = 140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, width in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(width))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def clear_cell(cell) -> None:
    cell.text = ""
    for para in cell.paragraphs:
        para.text = ""


def cell_para(
    cell,
    text: str,
    *,
    size: float = 9.2,
    color: str = INK,
    bold: bool = False,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
    before: float = 0,
    after: float = 2,
    line_spacing: float = 1.08,
):
    para = cell.add_paragraph() if cell.paragraphs and cell.paragraphs[0].text else cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = line_spacing
    run = para.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return para


def prepare_cell(cell, fill: str | None = None) -> None:
    if fill:
        set_cell_fill(cell, fill)
    set_cell_margins(cell)
    set_cell_border(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def add_data_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_cm: list[float],
    *,
    header_fill: str = NAVY,
    body_fill: str = WHITE,
    zebra: bool = True,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, widths_cm)

    for idx, title in enumerate(headers):
        cell = table.rows[0].cells[idx]
        clear_cell(cell)
        prepare_cell(cell, header_fill)
        cell_para(cell, title, size=8.8, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            fill = "F8FAFB" if zebra and row_idx % 2 else body_fill
            clear_cell(cells[col_idx])
            prepare_cell(cells[col_idx], fill)
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx in (0, 1) and len(value) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
            cell_para(cells[col_idx], value, size=8.7, color=INK, align=align)

    para_text(doc, "", after=4)
    return table


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    # Chinese formal reports are normally reviewed in A4; this is the named layout override.
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_EN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.5)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def setup_header_footer(doc: Document, company_name: str, report_title: str = "企业碳资信评级报告") -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(3)
    run = header.add_run(f"{report_title} | {company_name}")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    paragraph_border_bottom(header, BORDER, "6")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("模型输出供客户审核确认 | 第 ")
    set_run_font(run, size=8, color=MUTED)
    add_page_number(footer)
    run = footer.add_run(" 页")
    set_run_font(run, size=8, color=MUTED)


def grade_fill(grade: str | None, no_rating: bool) -> tuple[str, str]:
    if no_rating:
        return PALE_RED, RED
    if grade in {"AAA", "AA", "A"}:
        return PALE_GREEN, GREEN
    if grade in {"BBB", "BB", "B"}:
        return PALE_GOLD, GOLD
    return PALE_RED, RED


def reason_text(reasons: list[dict[str, Any]]) -> str:
    if not reasons:
        return ""
    parts = []
    for reason in reasons:
        r_type = reason.get("type", "不评级事项")
        fields = reason.get("fields")
        if fields:
            parts.append(f"{r_type}: {', '.join(fields)}")
        else:
            parts.append(r_type)
    return "；".join(parts)


def normalize_report_result(result: dict[str, Any]) -> dict[str, Any]:
    """Normalize nested scorer output and manually curated rating JSON."""

    result = dict(result)
    source = result.get("source_payload") or {}
    for key in [
        "report_title",
        "report_number",
        "report_date",
        "veto_check",
        "basic_credit_detail",
        "swv_details",
        "transition_detail",
        "report_statement",
        "recommendations",
        "risk_prompts",
        "agency",
        "missing_fields",
    ]:
        if key not in result and key in source:
            result[key] = source[key]
    enterprise = dict(result.get("enterprise") or {})
    if not enterprise:
        for key in [
            "company_name",
            "unified_credit_code",
            "industry",
            "industry_class",
            "rating_period",
            "registered_address",
            "report_date",
            "report_number",
            "development_stage",
            "carbon_account_filing_no",
        ]:
            if key in source:
                enterprise[key] = source[key]
            elif key in result:
                enterprise[key] = result[key]
    for key in ["industry_class", "report_date", "report_number", "development_stage", "carbon_account_filing_no"]:
        if key in source and key not in enterprise:
            enterprise[key] = source[key]
        if key in result and key not in enterprise:
            enterprise[key] = result[key]

    scores = dict(result.get("scores") or {})
    for key in ["basic_credit_score", "carbon_account_score", "transition_score", "comprehensive_score", "grade"]:
        if key in source and key not in scores:
            scores[key] = source[key]
        if key in result and key not in scores:
            scores[key] = result[key]

    weights = dict(result.get("weights") or {})
    if {"alpha", "beta", "gamma"}.issubset(weights):
        weights = {
            "abc": {"alpha": weights["alpha"], "beta": weights["beta"], "gamma": weights["gamma"]},
            "swv": {"s": weights.get("Ws"), "w": weights.get("Ww"), "v": weights.get("Wv")},
            **weights,
        }
    if "abc" not in weights and "weights" in source:
        src_weights = source["weights"]
        if {"alpha", "beta", "gamma"}.issubset(src_weights):
            weights["abc"] = {"alpha": src_weights["alpha"], "beta": src_weights["beta"], "gamma": src_weights["gamma"]}
        if {"Ws", "Ww", "Wv"}.issubset(src_weights):
            weights["swv"] = {"s": src_weights["Ws"], "w": src_weights["Ww"], "v": src_weights["Wv"]}
    result["enterprise"] = enterprise
    result["scores"] = scores
    result["weights"] = weights
    if "industry_class" not in result:
        result["industry_class"] = enterprise.get("industry_class") or source.get("industry_class")
    return result


def display_industry_class(result: dict[str, Any]) -> str:
    return fmt(result.get("industry_class") or result.get("enterprise", {}).get("industry_class"))


def add_section_title(doc: Document, number: str, title: str, subtitle: str | None = None) -> None:
    para = para_text(doc, "", before=10, after=3)
    badge = para.add_run(number)
    set_run_font(badge, size=8.5, color=TEAL, bold=True)
    title_run = para.add_run(f"  {title}")
    set_run_font(title_run, size=14, color=NAVY, bold=True)
    paragraph_border_bottom(para, BORDER, "6")
    if subtitle:
        para_text(doc, subtitle, size=9.2, color=MUTED, after=6)


def add_cover(doc: Document, result: dict[str, Any]) -> None:
    enterprise = result.get("enterprise", {})
    scores = result.get("scores", {})
    no_rating = bool(result.get("no_rating"))
    grade = None if no_rating else scores.get("grade")
    fill, grade_color = grade_fill(grade, no_rating)

    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band.style = "Table Grid"
    set_table_widths(band, [17.2])
    cell = band.rows[0].cells[0]
    clear_cell(cell)
    prepare_cell(cell, NAVY)
    cell_para(cell, "CARBON CREDIT RATING REPORT", size=8.5, color="B9D6D4", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    para_text(doc, fmt(result.get("report_title", "企业碳资信评级报告")), size=25, color=NAVY, bold=True, before=22, after=4)
    para_text(doc, fmt(enterprise.get("company_name")), size=14, color=SLATE, bold=True, after=14)

    scorecard = doc.add_table(rows=1, cols=2)
    scorecard.alignment = WD_TABLE_ALIGNMENT.CENTER
    scorecard.style = "Table Grid"
    set_table_widths(scorecard, [5.0, 12.2])

    left, right = scorecard.rows[0].cells
    clear_cell(left)
    prepare_cell(left, fill)
    cell_para(left, "评级结果", size=9, color=grade_color, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    cell_para(left, "不评级" if no_rating else fmt(grade), size=28, color=grade_color, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    cell_para(left, f"综合得分 {fmt(scores.get('comprehensive_score'))}", size=9, color=SLATE, align=WD_ALIGN_PARAGRAPH.CENTER)

    clear_cell(right)
    prepare_cell(right, MIST)
    conclusion = reason_text(result.get("no_rating_reasons", [])) if no_rating else "未触发直接不评级事项，按综合得分输出等级。"
    cell_para(right, "评级摘要", size=10, color=NAVY, bold=True, after=4)
    cell_para(right, conclusion, size=9.2, color=INK, after=6)
    cell_para(
        right,
        f"行业分类：{display_industry_class(result)}    评价周期：{fmt(enterprise.get('rating_period'))}",
        size=8.7,
        color=MUTED,
    )

    para_text(doc, "", after=8)
    add_metric_strip(
        doc,
        [
            ("基本信用", fmt(scores.get("basic_credit_score"))),
            ("碳账户", fmt(scores.get("carbon_account_score"))),
            ("绿色转型", fmt(scores.get("transition_score"))),
            ("评级状态", "不评级" if no_rating else "已评级"),
        ],
    )

    meta_rows = [
        ["统一社会信用代码", fmt(enterprise.get("unified_credit_code")), "所属行业", fmt(enterprise.get("industry"))],
        ["行业分类", display_industry_class(result), "报告日期", fmt(enterprise.get("report_date") or result.get("report_date") or date.today().isoformat())],
        ["报告编号", fmt(enterprise.get("report_number") or result.get("report_number")), "碳账户备案号", fmt(enterprise.get("carbon_account_filing_no") or result.get("carbon_account_filing_no"))],
        ["发展阶段", fmt(enterprise.get("development_stage") or result.get("development_stage")), "严重失信", "是" if result.get("veto_check", {}).get("serious_dishonesty_flag") else "否"],
        ["模型版本", fmt(result.get("source_payload", {}).get("model_version", "enterprise-carbon-credit-rating v0.1")), "参数来源", "行业默认参数，生成前由用户确认"],
    ]
    add_data_table(doc, ["项目", "内容", "项目", "内容"], meta_rows, [3.5, 5.1, 3.5, 5.1], header_fill=TEAL, zebra=False)
    para_text(doc, "说明：本报告为企业碳资信评价模型输出结果，需结合客户确认的数据、行业参数及原始佐证材料使用。", size=8.6, color=MUTED, italic=True, before=10, after=0)
    para_text(doc, "", after=8)


def add_metric_strip(doc: Document, items: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, [17.2 / len(items)] * len(items))
    for idx, (label, value) in enumerate(items):
        cell = table.rows[0].cells[idx]
        clear_cell(cell)
        prepare_cell(cell, PALE_BLUE if idx % 2 == 0 else MIST)
        cell_para(cell, label, size=8.3, color=MUTED, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
        cell_para(cell, value, size=15, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
    para_text(doc, "", after=6)


def add_grade_scale(doc: Document, current_grade: str | None, no_rating: bool) -> None:
    bands = [
        ("AAA", "≥90"),
        ("AA", "85-90"),
        ("A", "80-85"),
        ("BBB", "75-80"),
        ("BB", "70-75"),
        ("B", "65-70"),
        ("CCC", "60-65"),
        ("CC", "50-60"),
        ("C", "<50"),
    ]
    table = doc.add_table(rows=1, cols=len(bands))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, [17.2 / len(bands)] * len(bands))
    for idx, (grade, band) in enumerate(bands):
        cell = table.rows[0].cells[idx]
        clear_cell(cell)
        is_current = not no_rating and grade == current_grade
        prepare_cell(cell, PALE_GREEN if is_current else WHITE)
        cell_para(cell, grade, size=8.5, color=GREEN if is_current else SLATE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
        cell_para(cell, band, size=7.2, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    para_text(doc, "", after=4)


def add_weight_block(doc: Document, result: dict[str, Any]) -> None:
    weights = result.get("weights", {})
    abc = weights.get("abc", {})
    swv = weights.get("swv") or {}
    rows = [
        ["α 基本信用", pct(abc.get("alpha")), "综合评分权重", "客户指南依据；行业配置"],
        ["β 碳账户", pct(abc.get("beta")), "综合评分权重", "客户指南依据；行业配置"],
        ["γ 绿色低碳转型", pct(abc.get("gamma")), "综合评分权重", "客户指南依据；行业配置"],
        ["S 碳资产价值", pct(swv.get("s")), "碳账户内部权重", "客户指南依据；行业配置"],
        ["W 碳强度对标", pct(swv.get("w")), "碳账户内部权重", "客户指南依据；行业配置"],
        ["V 碳行为", pct(swv.get("v")), "碳账户内部权重", "客户指南依据；行业配置"],
    ]
    add_data_table(doc, ["权重项", "取值", "用途", "来源标注"], rows, [4.0, 2.3, 4.2, 6.7], header_fill=NAVY)


def add_component_scores(doc: Document, result: dict[str, Any]) -> None:
    scores = result.get("scores", {})
    rows = [
        ["基本信用", fmt(scores.get("basic_credit_score")), "信用合规、财务能力、生产与合规管理、创新与品牌建设加权。", "客户指南依据"],
        ["碳账户 SWV", fmt(scores.get("carbon_account_score")), "S 反映碳资产价值，W 反映碳强度相对行业水平，V 反映碳行为异常。", "客户指南依据"],
        ["绿色低碳转型", fmt(scores.get("transition_score")), "自愿减排项目、转型绩效、技改降碳实施成效加权。", "客户指南依据"],
    ]
    add_data_table(doc, ["维度", "得分", "评价逻辑", "来源标注"], rows, [3.2, 2.2, 8.6, 3.2], header_fill=TEAL)

    detail = result.get("basic_credit_detail", {})
    if detail.get("dimension_rows"):
        add_data_table(
            doc,
            ["维度", "权重", "得分", "关键表现概述"],
            [[fmt(x.get("name")), fmt(x.get("weight")), fmt(x.get("score")), fmt(x.get("summary"))] for x in detail["dimension_rows"]],
            [3.2, 2.0, 2.2, 9.8],
            header_fill=SLATE,
        )


def add_swv_details(doc: Document, result: dict[str, Any]) -> None:
    swv_details = result.get("swv_details", {})
    if not swv_details:
        return
    component = swv_details.get("component_scores", {})
    headers = ["项目", "分值", "测算值", "来源标注"]
    if swv_details.get("dimension_rows"):
        rows = [[fmt(x.get("name")), fmt(x.get("score")), fmt(x.get("value")), fmt(x.get("summary"))] for x in swv_details["dimension_rows"]]
        headers = ["项目", "分值", "测算值", "关键指标值"]
    else:
        rows = [
            ["S 碳资产价值", fmt(component.get("s")), fmt(swv_details.get("s", {}).get("value")), "客户指南依据"],
            ["W 碳强度对标", fmt(component.get("w")), fmt(swv_details.get("w", {}).get("value")), "客户指南依据"],
            ["V 碳行为", fmt(component.get("v")), fmt(swv_details.get("v", {}).get("value")), "客户指南依据"],
        ]
    add_data_table(doc, headers, rows, [4.0, 2.5, 5.3, 5.4], header_fill=TEAL)


def add_transition_details(doc: Document, result: dict[str, Any]) -> None:
    detail = result.get("transition_detail", {})
    if not detail.get("dimension_rows"):
        return
    add_data_table(
        doc,
        ["维度", "权重", "得分", "关键指标值"],
        [[fmt(x.get("name")), fmt(x.get("weight")), fmt(x.get("score")), fmt(x.get("summary"))] for x in detail["dimension_rows"]],
        [3.5, 2.0, 2.2, 9.5],
        header_fill=TEAL,
    )


def add_adjustment_matrix(doc: Document, result: dict[str, Any]) -> None:
    no_rating = bool(result.get("no_rating"))
    reasons = reason_text(result.get("no_rating_reasons", []))
    fill = PALE_RED if no_rating else PALE_GREEN
    color = RED if no_rating else GREEN
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, [17.2])
    cell = table.rows[0].cells[0]
    clear_cell(cell)
    prepare_cell(cell, fill)
    cell_para(cell, "不评级判断", size=9, color=color, bold=True, after=2)
    cell_para(cell, f"本期触发事项：{reasons}" if no_rating else "本期未触发直接不评级事项。", size=9.5, color=INK, after=0)
    para_text(doc, "", after=4)

    rows = [
        ["数据造假", "直接不评级", "客户确认规则"],
        ["重大环境处罚", "直接不评级", "客户确认规则；模型建议纳入风险底线"],
        ["重大银行不良记录", "直接不评级", "客户确认规则；模型建议纳入风险底线"],
        ["严重失信主体名单", "直接不评级", "客户指南依据"],
        ["综合得分低于 40", "输出 C，不自动不评级", "客户确认规则"],
    ]
    add_data_table(doc, ["事项", "处理规则", "来源标注"], rows, [4.5, 5.0, 7.7], header_fill=NAVY)


def add_missing_data(doc: Document, result: dict[str, Any]) -> None:
    source_payload = result.get("source_payload", {})
    missing = result.get("missing", {})
    rows = []
    for field in missing.get("identity_required", []):
        rows.append([field, "不可缺失基础字段", "需补全后方可评分"])
    for field in missing.get("key_required", []):
        rows.append([field, "关键必要指标", "用户确认保持缺失时输出不评级"])
    for field in source_payload.get("missing_noncritical", []):
        rows.append([field, "非关键指标", "披露为未采集数据"])
    for field in result.get("missing_fields", []):
        rows.append([field, "待补充/待核验", "报告中披露，不使用默认分替代"])
    if not rows:
        rows.append(["无", "无", "本次样例未识别到需披露的缺失字段"])
    add_data_table(doc, ["字段", "类别", "处理方式"], rows, [5.0, 4.4, 7.8], header_fill=SLATE)


def add_risks_and_recommendations(doc: Document, result: dict[str, Any]) -> None:
    source_payload = result.get("source_payload", {})
    risks = result.get("risk_prompts") or source_payload.get("risk_prompts") or ["需持续关注碳数据质量、环境合规记录、银行信用记录及行业参数适配性。"]
    recommendations = result.get("recommendations") or source_payload.get("recommendations") or ["建议完善碳排放核算、减排项目台账、绿色电力凭证和技改降碳成效证明材料。"]
    rows = []
    for idx, item in enumerate(risks, 1):
        rows.append([f"R{idx}", "风险提示", item])
    for idx, item in enumerate(recommendations, 1):
        rows.append([f"A{idx}", "改进建议", item])
    add_data_table(doc, ["编号", "类型", "内容"], rows, [2.0, 3.0, 12.2], header_fill=NAVY)


def add_report_statement(doc: Document, result: dict[str, Any]) -> None:
    statement = result.get("report_statement") or [
        "本报告依据《企业碳征信报告编制指南》及客户确认的数据口径编制。",
        "评价机构承诺独立、客观、公正地开展评估；重大指标经第三方核证或客户确认后纳入。",
        "本报告仅反映评价周期内企业碳资信状况，数据来源详见附件或客户提供资料。",
    ]
    rows = [[str(idx), item] for idx, item in enumerate(statement, 1)]
    add_data_table(doc, ["序号", "声明内容"], rows, [2.0, 15.2], header_fill=SLATE)


def add_agency_info(doc: Document, result: dict[str, Any]) -> None:
    agency = result.get("agency") or {}
    if not agency:
        return
    rows = [
        ["机构名称", fmt(agency.get("name")), "统一社会信用代码", fmt(agency.get("unified_credit_code"))],
        ["地址", fmt(agency.get("address")), "联系电话", fmt(agency.get("phone"))],
        ["资质说明", fmt(agency.get("qualification")), "备注", fmt(agency.get("note"))],
    ]
    add_data_table(doc, ["项目", "内容", "项目", "内容"], rows, [2.8, 5.8, 3.2, 5.4], header_fill=SLATE, zebra=False)


def build_doc(result: dict[str, Any], out_path: Path) -> None:
    result = normalize_report_result(result)
    doc = Document()
    setup_document(doc)

    enterprise = result.get("enterprise", {})
    scores = result.get("scores", {})
    no_rating = bool(result.get("no_rating"))
    grade = None if no_rating else scores.get("grade")
    no_rating_detail = reason_text(result.get("no_rating_reasons", []))
    setup_header_footer(doc, fmt(enterprise.get("company_name")), fmt(result.get("report_title", "企业碳资信评级报告")))

    add_cover(doc, result)

    add_section_title(doc, "01", "评级结论", "参考评级机构报告惯例，将评级结果、综合得分、行业参数和直接不评级事项前置披露。")
    conclusion = (
        f"因{no_rating_detail}，本期不出具企业碳资信等级。"
        if no_rating
        else f"本期综合得分为 {fmt(scores.get('comprehensive_score'))} 分，评级结果为 {fmt(grade)}。"
    )
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, [17.2])
    cell = table.rows[0].cells[0]
    clear_cell(cell)
    fill, color = grade_fill(grade, no_rating)
    prepare_cell(cell, fill)
    cell_para(cell, "核心结论", size=9, color=color, bold=True, after=2)
    cell_para(cell, conclusion, size=11, color=INK, bold=True, after=2)
    cell_para(cell, "直接不评级事项优先于分数等级；未触发时按综合得分区间输出等级。", size=8.7, color=MUTED, after=0)
    para_text(doc, "", after=4)

    add_metric_strip(
        doc,
        [
            ("综合等级", "不评级" if no_rating else fmt(grade)),
            ("综合得分", fmt(scores.get("comprehensive_score"))),
            ("行业分类", display_industry_class(result)),
            ("评价周期", fmt(enterprise.get("rating_period"))),
        ],
    )
    add_grade_scale(doc, grade, no_rating)

    add_section_title(doc, "02", "报告声明")
    add_report_statement(doc, result)

    add_section_title(doc, "03", "企业基础信息")
    add_data_table(
        doc,
        ["项目", "内容", "项目", "内容"],
        [
            ["企业名称", fmt(enterprise.get("company_name")), "统一社会信用代码", fmt(enterprise.get("unified_credit_code"))],
            ["注册地址", fmt(enterprise.get("registered_address")), "所属行业", fmt(enterprise.get("industry"))],
            ["评价周期", fmt(enterprise.get("rating_period")), "发展阶段", fmt(enterprise.get("development_stage"))],
            ["碳账户备案号", fmt(enterprise.get("carbon_account_filing_no")), "资料来源", fmt(result.get("source_payload", {}).get("data_source", "客户提供资料及模型抽取结果"))],
        ],
        [3.5, 5.1, 3.5, 5.1],
        header_fill=TEAL,
        zebra=False,
    )

    add_section_title(doc, "04", "模型方法与权重", "权重和行业分类在正式生成报告前需由用户审核确认；污水处理行业采用客户样例报告特例。")
    para_text(doc, "综合评分 = 基本信用得分 × α + 碳账户得分 × β + 绿色低碳转型得分 × γ。碳账户得分由 S、W、V 三项加权形成。", size=9.5, color=INK, after=6)
    add_weight_block(doc, result)

    add_section_title(doc, "05", "分项评价结果")
    add_component_scores(doc, result)
    add_swv_details(doc, result)
    add_transition_details(doc, result)

    doc.add_page_break()
    add_section_title(doc, "06", "调整项与不评级判断")
    add_adjustment_matrix(doc, result)

    add_section_title(doc, "07", "缺失数据披露")
    add_missing_data(doc, result)

    add_section_title(doc, "08", "风险提示与改进建议")
    add_risks_and_recommendations(doc, result)

    add_section_title(doc, "09", "评价机构信息")
    add_agency_info(doc, result)

    add_section_title(doc, "10", "附录：来源标注说明")
    add_data_table(
        doc,
        ["标注", "含义"],
        [
            ["客户指南依据", "来自客户提供的《基于企业碳账户的碳资信评价标准》"],
            ["客户样例/确认", "来自客户样例报告或本次对话中确认的业务规则"],
            ["模型建议", "为提高评级稳健性而补充的风险提示或调整项"],
            ["外部惯例补充", "参照温室气体核算、可持续信息披露及金融机构碳核算常见实践"],
        ],
        [4.0, 13.2],
        header_fill=SLATE,
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("rating_result")
    parser.add_argument("--out", default="企业碳资信评级报告.docx")
    args = parser.parse_args()

    result = json.loads(Path(args.rating_result).read_text(encoding="utf-8"))
    out_path = Path(args.out)
    build_doc(result, out_path)
    print(out_path)


if __name__ == "__main__":
    main()
